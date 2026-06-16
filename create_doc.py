import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Styles
title = doc.add_heading('Refract Platform Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1: Overview
heading = doc.add_heading('1. Project Overview', level=1)
p = doc.add_paragraph('Refract (formerly "Lumina") is an open-source AI observability, telemetry, and cost-optimization platform. Built for the AiGENThix Hackathon, it serves as a transparent middleman between client applications and large language models (LLMs) such as Google Gemini, OpenAI, and Anthropic. Refract captures traces, calculates costs, analyzes prompt efficiency, and provides an extensive metrics dashboard to monitor AI usage without requiring significant code changes in client applications.')

# Section 2: Core Architecture & Components
doc.add_heading('2. Core Architecture & Implemented Features', level=1)
p = doc.add_paragraph('The platform was transformed from a basic observability clone into a robust proxy-based system through several major implementation phases:')

p = doc.add_paragraph('Phase 0: Rebranding (Lumina → Refract)', style='List Bullet')
p.add_run('\nExecuted a global rename across the entire codebase, including package names, environment variables, docker configurations, SDKs, database schemas, and all user-facing UI elements. We also removed legacy image assets and replaced them with a sleek, text-based SVG logo with dynamic gradients and styling.')

p = doc.add_paragraph('Phase 1: Database Schema Updates', style='List Bullet')
p.add_run('\nExpanded the PostgreSQL schema to support advanced prompt analysis. Added columns for prompt category, complexity, model fit, source (SDK vs Proxy), and updated the provider constraints to fully support Google Gemini alongside OpenAI and Anthropic.')

p = doc.add_paragraph('Phase 2: Proxy Gateway Service', style='List Bullet')
p.add_run('\nBuilt a transparent HTTP middleman service (port 8090) that intercepts LLM calls. It extracts prompt and model metadata, forwards requests to the real LLM provider, and asynchronously sends telemetry to the Ingestion service (fire-and-forget) without adding latency to the client. This allows any application to be instrumented simply by changing its base URL to point to the proxy.')

p = doc.add_paragraph('Phase 3: Prompt Analyzer Engine', style='List Bullet')
p.add_run('\nImplemented an on-demand prompt analysis engine. When triggered, it uses a lightweight LLM call to categorize prompts, rate their complexity, and assess if the chosen model is overkill, underkill, or a good fit for the task. This engine directly calculates potential cost savings (Model Cascading Opportunities).')

p = doc.add_paragraph('Phase 4: Demo Chat Interface', style='List Bullet')
p.add_run('\nCreated a built-in Demo Chat UI that automatically routes through the Refract Proxy. It allows users to test different models (e.g., gemini-2.0-flash, gemini-2.5-pro) and generates real traces, metrics, latency, and token cost data instantly visible on the dashboard.')

p = doc.add_paragraph('Phase 5: Dashboard Enhancements & Unique KPIs', style='List Bullet')
p.add_run('\nDeveloped a suite of Refract-exclusive visualizations to differentiate the platform:')
p.add_run('\n- Prompt Efficiency Score: Ratio of output to input tokens.')
p.add_run('\n- Token Flow Breakdown: Stacked bar visualization of input vs output.')
p.add_run('\n- Model-Task Alignment: Distribution chart of overkill vs good fit.')
p.add_run('\n- Estimated Savings: Dollar-value metric for potential cost optimization.')

# Section 3: Recent Fixes & Stability
doc.add_heading('3. Key Bug Fixes & Stability Improvements', level=1)

p = doc.add_paragraph('1. Database Ingestion Fix:\n', style='List Bullet')
p.add_run('Resolved a critical constraint violation where OpenTelemetry traces were dropping data. Implemented a data transformation layer to seamlessly map snake_case telemetry to the camelCase expectations of the Drizzle ORM, ensuring successful data persistence.')

p = doc.add_paragraph('2. API Connectivity & Service Routing:\n', style='List Bullet')
p.add_run('Fixed "Failed to fetch" errors on the Replay and Traces pages. Corrected environment configurations and routing logic to properly resolve the Query API (port 8081) and Replay Engine (port 8082).')

p = doc.add_paragraph('3. Prompt Analysis Rate Limit Handling:\n', style='List Bullet')
p.add_run('Fixed an authentication customerID mismatch that hid traces from the analysis engine. Additionally, implemented a robust fallback strategy for Gemini API 429 RESOURCE_EXHAUSTED errors. When the free-tier quota is reached, the system injects mock analysis data to ensure the dashboard visualizations continue functioning during demonstrations.')

p = doc.add_paragraph('4. Trace Deletion Functionality:\n', style='List Bullet')
p.add_run('Upgraded the UI placeholder alert to a fully functional DELETE endpoint. Users can now securely delete traces (which cascades to associated alerts) and the UI updates instantaneously.')

p = doc.add_paragraph('5. UI State Mapping Correction:\n', style='List Bullet')
p.add_run('Corrected the React trace mapping logic. Successfully processed proxy traces mapped as "success" were being misinterpreted as errors; they are now correctly displayed as "healthy" traces in the dashboard table.')

# Section 4: External Integrations
doc.add_heading('4. Integrating External Applications', level=1)
p = doc.add_paragraph('Refract fully supports monitoring external applications. Developers can connect their AI projects using two distinct methods:')

p = doc.add_paragraph()
p.add_run('Method A: Transparent LLM Proxy (Zero-Code Integration)\n').bold = True
p.add_run('External applications can route their OpenAI or Gemini API calls through the Refract Proxy (http://localhost:8090/v1) by modifying their SDK\'s base URL. The proxy seamlessly forwards the request, manages API keys, and logs all telemetry automatically.')

p = doc.add_paragraph()
p.add_run('Method B: Refract SDK Integration\n').bold = True
p.add_run('For applications requiring deep instrumentation, developers can install the Refract SDK (@refract/sdk) to directly track custom functions, wrap logic, and send customized telemetry data directly to the Ingestion service.')

# Section 5: Installation and Running
doc.add_heading('5. Installation & Running Guide', level=1)

p = doc.add_paragraph('1. Prerequisites:\n', style='List Bullet')
p.add_run('Ensure you have Docker Desktop, Node.js, and Bun installed.')

p = doc.add_paragraph('2. Install Dependencies:\n', style='List Bullet')
p.add_run('Run the following in the project root:\n')
p.add_run('bun install').font.name = 'Courier New'

p = doc.add_paragraph('3. Start Infrastructure:\n', style='List Bullet')
p.add_run('Launch PostgreSQL, NATS, and Redis via Docker:\n')
p.add_run('docker compose -f infra/docker/docker-compose.yml up -d').font.name = 'Courier New'

p = doc.add_paragraph('4. Database Migrations:\n', style='List Bullet')
p.add_run('Initialize the schema:\n')
p.add_run('bun run db:push').font.name = 'Courier New'

p = doc.add_paragraph('5. Start Microservices:\n', style='List Bullet')
p.add_run('In separate terminal windows, launch the backend services:\n')
p.add_run('- bun run dev:ingestion (Port 8080)\n').font.name = 'Courier New'
p.add_run('- bun run dev:api (Port 8081)\n').font.name = 'Courier New'
p.add_run('- bun run dev:replay (Port 8082)\n').font.name = 'Courier New'
p.add_run('- bun run dev:proxy (Port 8090)').font.name = 'Courier New'

p = doc.add_paragraph('6. Start the Dashboard UI:\n', style='List Bullet')
p.add_run('Launch the frontend:\n')
p.add_run('bun run dev').font.name = 'Courier New'
p.add_run('\nThe dashboard is available at http://localhost:3000, and the Demo Chat UI can be accessed at http://localhost:8090/demo.')

# Save document
doc_path = os.path.join(os.getcwd(), 'Refract_Project_Documentation_Complete.docx')
doc.save(doc_path)
print(f"Document saved successfully to {doc_path}")
